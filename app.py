# -*- coding: utf-8 -*-
"""
MEDIBOT - Medical Domain RAG Chatbot
With Fixed Memory + Strict Medical Q-Type + Semantic Evaluation
Includes:
    - Drug Info Agent (via OpenFDA API)
    - Body Metrics (BMI) Agent (Manual Calculation Only)
    - Diagnosis Agent (via PubMed + RAG reasoning + smart symptom trigger)
    - Lifestyle & Prevention Agent (WGER + LLM hybrid)
    - Image Agent (Gemini primary + Hugging Face Stable Diffusion XL backup)
LLM fallback enabled for questions with no FAISS context
Created: 2025-10-28
@author: Arun
"""

# ============================ #
# Step 0: Imports
# ============================ #
from langchain.chains import ConversationalRetrievalChain
from langchain.prompts import PromptTemplate
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAI
from langchain.memory import ConversationBufferMemory
from sklearn.metrics.pairwise import cosine_similarity
from langchain.docstore.document import Document
import re, requests, os
from dotenv import load_dotenv
import base64
from io import BytesIO
from PIL import Image
from datetime import datetime
import matplotlib.pyplot as plt
import google.generativeai as genai
from huggingface_hub import InferenceClient
import streamlit as st
import time

# ============================ #
# Step 1: Load Environment Keys
# ============================ #
ENV_PATH = r"D:\AAIDC\Project 2\.env"
load_dotenv(ENV_PATH)
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
WGER_API_KEY = os.getenv("WGER_API_KEY")
HF_TOKEN = os.getenv("HF_TOKEN")

# ============================ #
# Step 2: Load FAISS DB + Embeddings
# ============================ #
DB_FAISS_PATH = r"D:\AAIDC\Project 2\vectorstore"
embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2",
    model_kwargs={"device": "cpu"}  # ✅ Force CPU, fixes meta tensor issue
)
db = FAISS.load_local(DB_FAISS_PATH, embeddings, allow_dangerous_deserialization=True)
retriever = db.as_retriever(search_kwargs={"k": 10})

# ============================ #
# Step 3: Gemini LLM
# ============================ #
llm = GoogleGenerativeAI(
    model="gemini-2.0-flash",
    google_api_key=GOOGLE_API_KEY,
    temperature=0.2,
    max_output_tokens=1200
)

# ============================ #
# Step 4: Prompt Template
# ============================ #
FULL_PROMPT_TEMPLATE = """
You are MEDIBOT, a strict AI medical tutor and consultant.
Answer only medical-related questions based on the provided context if available.
If context is empty, fallback to general medical knowledge:
- Include recent studies, clinical trials, or research findings (up to your knowledge cutoff)
- Include clinical relevance, examples, and treatment implications
- Give approximate references or citations in readable format

Do NOT provide answers outside the medical domain.

---

### Memory & Context Rules:
- Use chat history to interpret vague follow-ups
- Continue the flow instead of repeating explanations
- Only answer medical questions; refuse irrelevant queries
- Always use **chat history** to interpret vague follow-ups (e.g., "it", "this", "go with that").
- If the user gives acknowledgments like "okay", "yes", "continue", interpret them as **follow-up requests**.

---

### Depth Control:
- Expand answers into at least 4–5 lines
- Include key context, clinical reasoning, and examples
- Avoid any non-medical content

---

Chat History:
{chat_history}

Context from medical material (if any):
{context}

Patient / User Question:
{question}

Answer:
"""

FULL_PROMPT = PromptTemplate(
    template=FULL_PROMPT_TEMPLATE,
    input_variables=["chat_history", "context", "question"]
)

# ============================ #
# Step 5: Memory
# ============================ #
memory = ConversationBufferMemory(
    memory_key="chat_history",
    return_messages=True,
    output_key="answer"
)

summary_memory = ""

def update_summary_memory(chat_history, max_recent_turns=5):
    global summary_memory
    if len(chat_history) <= max_recent_turns:
        return ""
    old_msgs = chat_history[:-max_recent_turns]
    summary_memory = " ".join([f"{m.type}: {m.content}" for m in old_msgs])
    return summary_memory

def trim_chat_history(chat_history, max_recent_turns=5):
    return chat_history[-max_recent_turns:]

def get_combined_history(chat_history, max_recent_turns=5):
    trimmed = trim_chat_history(chat_history, max_recent_turns)
    combined = ""
    if summary_memory:
        combined += f"[Summary of older conversation]: {summary_memory}\n"
    for msg in trimmed:
        combined += f"{msg.type}: {msg.content}\n"
    return combined


# ============================ #
# Step 6: Conversational Retrieval Chain
# ============================ #
medical_chain = ConversationalRetrievalChain.from_llm(
    llm=llm,
    retriever=retriever,
    memory=memory,
    combine_docs_chain_kwargs={"prompt": FULL_PROMPT},
    return_source_documents=False,
    output_key="answer"
)

# ============================ #
# Step 7: Helpers
# ============================ #
def semantic_similarity_score(reference, generated, embed_model=embeddings):
    if not reference or not generated:
        return None
    ref_vec = embed_model.embed_query(reference)
    gen_vec = embed_model.embed_query(generated)
    score = cosine_similarity([ref_vec], [gen_vec])[0][0]
    return round(score, 4)

def evaluate_response(reference, generated):
    return {"SemanticSim": semantic_similarity_score(reference, generated)}


# ✅ Detect vague or follow-up questions
def detect_followup(query):
    """
    Detects follow-up intent like 'go with this', 'for it', 'the 1st one', 'can you give me more',
    or specific references like 'explain 2nd point' or 'go with 5th one'.
    """
    followup_patterns = [
        # General follow-up / continuation
        r"\bok\b", r"\bokay\b", r"\byes\b", r"\bcontinue\b", r"\bgo on\b", r"\bnext\b",
        r"\bwhat next\b", r"\bnext step\b", r"\bgo ahead\b", r"\bkeep going\b",
        r"\btell me more\b", r"\bmore about\b", r"\bmore on\b", r"\bexplain more\b",
        r"\belaborate\b", r"\bexpand\b", r"\bgo deeper\b", r"\bexplain further\b",
        r"\bcan you give me more\b", r"\bdetails\b",

        # Contextual pronouns
        r"\bfor it\b", r"\babout it\b", r"\brelated to it\b", r"\babout that\b",
        r"\bfor that\b", r"\bregarding that\b", r"\bon this\b", r"\bthis topic\b",
        r"\bthat one\b", r"\bthis one\b", r"\bthese\b", r"\bthose\b", r"\bit\b",

        # Numeric point references (1st, 2nd, third, etc.)
        r"\b(go with|explain|elaborate on|describe|tell me about|for|choose|select)\s*\d+(st|nd|rd|th)\b",
        r"\b(point|option|number)\s*\d+\b",
        r"\bthe\s*\d+(st|nd|rd|th)\s*(point|one|option)\b"
    ]

    pattern = re.compile("|".join(followup_patterns), re.IGNORECASE)
    return bool(re.search(pattern, query))



# 🔧 Updated follow-up detector
def is_medical_question(question):
    followups = r"\b(ok|okay|yes|continue|go with this|that one|explain\s*\d+|what about\s*\d+|the\s*\d+(st|nd|rd|th)\s*one|parts of it|functions of it|tell me more|details|more info)\b"
    if re.search(followups, question.lower()):
        return True
    non_medical_patterns = [
        r"\b(joke|funny|politics|movie|celebrity|personal)\b",
        r"\b(who|where|when) is .* president\b"
    ]
    for pat in non_medical_patterns:
        if re.search(pat, question.lower()):
            return False
    return True


# ============================ #
# Step 8: Agents
# ============================ #
OPENFDA_ENDPOINT = "https://api.fda.gov/drug/label.json"

def get_drug_info(drug_name):
    try:
        params = {"search": f"openfda.generic_name:{drug_name}", "limit": 1}
        response = requests.get(OPENFDA_ENDPOINT, params=params, timeout=10)
        if response.status_code != 200:
            return f"⚠️ Unable to fetch data from OpenFDA (Status: {response.status_code})"
        data = response.json()
        if "results" not in data or not data["results"]:
            return "⚠️ No official drug information found in OpenFDA."
        result = data["results"][0]
        openfda = result.get("openfda", {})
        brand = ", ".join(openfda.get("brand_name", ["N/A"]))
        generic = ", ".join(openfda.get("generic_name", ["N/A"]))
        indications = result.get("indications_and_usage", ["Not available"])[0]
        dosage = result.get("dosage_and_administration", ["Not available"])[0]
        warnings = result.get("warnings", ["Not available"])[0]
        contraindications = result.get("contraindications", ["Not available"])[0]
        return f"""💊 **Drug Information (via OpenFDA)**

**Brand & Generic Names:** {brand} | {generic}

**Indications:** {indications}

**Dosage:** {dosage}

**Warnings:** {warnings}

**Contraindications:** {contraindications}

📚 **Source: FDA Drug Label Database**"""
    except Exception as e:
        return f"❌ Error fetching drug info: {e}"

def is_drug_query(query):
    match = re.search(r"tell me about\s+([A-Za-z0-9\-]+)", query.lower())
    if match:
        return match.group(1)
    return None

def is_bmi_query(query):
    patterns = [r"\bbmi\b", r"\bbody\s*mass\b", r"\b(height|tall).*(weight|weigh)\b", r"\b\d+\s*(cm|m).*\d+\s*kg\b"]
    return any(re.search(p, query.lower()) for p in patterns)

def bmi_agent(query):
    query = query.lower()
    h_match = re.search(r"(\d+(\.\d+)?)\s*(cm|m)\b", query)
    w_match = re.search(r"(\d+(\.\d+)?)\s*kg\b", query)
    height = float(h_match.group(1)) / 100 if h_match and h_match.group(3) == "cm" else (float(h_match.group(1)) if h_match else None)
    weight = float(w_match.group(1)) if w_match else None
    if height and weight:
        bmi = round(weight / (height ** 2), 2)
        if bmi < 18.5:
            cat, adv = "Underweight", "Increase calorie intake with nutrient-rich foods."
        elif 18.5 <= bmi < 24.9:
            cat, adv = "Normal", "Maintain a balanced diet and regular exercise."
        elif 25 <= bmi < 29.9:
            cat, adv = "Overweight", "Exercise regularly and reduce processed food."
        else:
            cat, adv = "Obese", "Consult a healthcare provider for a structured plan."
        return f"🧍 **BMI Report**\n\n**Height:** {height:.2f} m\n**Weight:** {weight:.1f} kg\n**BMI:** {bmi}\n**Category:** {cat}\n💡 **Advice:** {adv}"
    return "Please provide height and weight clearly (e.g., 'height 170 cm weight 65 kg')."

def fetch_pubmed_articles(symptoms, max_results=5):
    try:
        base = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
        search_url = f"{base}esearch.fcgi"
        params = {"db": "pubmed", "term": symptoms, "retmax": max_results, "retmode": "json"}
        res = requests.get(search_url, params=params, timeout=10)
        data = res.json()
        ids = data["esearchresult"].get("idlist", [])
        if not ids:
            return []
        fetch_url = f"{base}efetch.fcgi"
        fetch_params = {"db": "pubmed", "id": ",".join(ids), "retmode": "text", "rettype": "abstract"}
        fetch_res = requests.get(fetch_url, params=fetch_params, timeout=10)
        return fetch_res.text.split("\n\n")[:max_results]
    except Exception as e:
        return [f"⚠️ Error fetching PubMed: {e}"]

def build_pubmed_vectorstore(abstracts):
    docs = [Document(page_content=a) for a in abstracts]
    db_temp = FAISS.from_documents(docs, embeddings)
    return db_temp.as_retriever(search_kwargs={"k": 5})

DIAGNOSIS_PROMPT = PromptTemplate(
    template="""You are MEDIBOT's Diagnosis Agent.
Given symptoms and PubMed data, identify possible diagnoses and reasoning.

Symptoms: {symptoms}
Context: {context}
Answer:""",
    input_variables=["symptoms", "context"]
)

def diagnosis_agent(symptoms):
    abstracts = fetch_pubmed_articles(symptoms)
    if not abstracts:
        return "⚠️ No PubMed results found for these symptoms."
    retriever_pubmed = build_pubmed_vectorstore(abstracts)
    docs = retriever_pubmed.get_relevant_documents(symptoms)
    context = " ".join([d.page_content for d in docs])
    prompt = DIAGNOSIS_PROMPT.format(symptoms=symptoms, context=context)
    try:
        resp = llm.invoke(prompt)
        return resp
    except Exception as e:
        return f"❌ Diagnosis generation failed: {e}"

def is_symptom_query(text):
    text = text.lower().strip()
    keywords = [
        "fever", "pain", "cough", "headache", "nausea", "vomit", "vomiting",
        "dizziness", "sore", "infection", "swelling", "fatigue", "cramps",
        "rash", "itch", "chills", "throat", "burning", "bleeding", "breath",
        "tightness", "diarrhea", "loss of taste", "loss of smell", "palpitation",
        "tremor", "sensitivity", "inflammation", "tingling", "sweating", "faint",
        "weakness", "pressure", "appetite", "stiffness", "cramp", "ache",
        "painful", "dizzy", "sick", "tired", "hurts"
    ]
    patterns = [
        r"\bi have\b", r"\bi’m having\b", r"\bim having\b", r"\bi am having\b",
        r"\bi feel\b", r"\bi’m feeling\b", r"\bim feeling\b", r"\bi am feeling\b",
        r"\bi got\b", r"\bi’ve got\b", r"\bi am suffering\b", r"\bi’m suffering\b",
        r"\bim suffering\b", r"\bi’ve been having\b", r"\bfeeling\b", r"\bhaving\b",
        r"\bsuffering from\b", r"\bmy\b"
    ]
    return any(k in text for k in keywords) and any(re.search(p, text) for p in patterns)

def lifestyle_agent(query, llm):
    llm_response = llm.invoke(f"You are LIFEGEN, an AI lifestyle coach. Give helpful advice for: {query}")
    base_answer = getattr(llm_response, "content", str(llm_response))
    return f"🧘 **Lifestyle Agent:**\n\n{base_answer}"

def research_agent(query, llm=None, memory=None):
    base_url = "https://www.ebi.ac.uk/europepmc/webservices/rest/search"
    params = {"query": query, "format": "json", "pageSize": 5}
    response = requests.get(base_url, params=params)
    if response.status_code != 200:
        return "❌ Failed to fetch research data."
    data = response.json()
    papers = data.get("resultList", {}).get("result", [])
    if not papers:
        return "⚕️ No research papers found."

    res = "🔬 **Top Medical Research Results:**\n\n"
    for i, paper in enumerate(papers, 1):
        title = paper.get("title", "No title")
        authors = paper.get("authorString", "Unknown authors")
        source = paper.get("source", "MED")
        paper_id = paper.get("id", "")
        link = f"https://europepmc.org/article/{source}/{paper_id}"
        res += (
            f"**{i}. {title}**\n"
            f"👨‍⚕️ *{authors}*\n"
            f"🔗 <a href='{link}' target='_blank'>Read Full Paper</a>\n\n"
        )
    return res



# ============================ #
# Step 9: Image Agent (Final - Shows Inline in Chat)
# ============================ #
if GOOGLE_API_KEY:
    genai.configure(api_key=GOOGLE_API_KEY)

hf_client = InferenceClient("stabilityai/stable-diffusion-xl-base-1.0", token=HF_TOKEN)
IMAGES_ROOT = os.path.join(os.path.dirname(__file__), "images")

def image_agent(prompt):
    """Generate image using Gemini or Hugging Face, save it, and show inline (base64-embedded)."""
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    folder = os.path.join(IMAGES_ROOT, ts)
    os.makedirs(folder, exist_ok=True)
    saved_path = None
    img = None

    # --- Try Gemini ---
    try:
        model = genai.GenerativeModel("gemini-2.5-flash-preview-image")
        resp = model.generate_content(prompt)
        for cand in resp.candidates:
            for part in cand.content.parts:
                inline = getattr(part, "inline_data", None)
                if inline and getattr(inline, "data", None):
                    img = Image.open(BytesIO(base64.b64decode(inline.data))).convert("RGBA")
                    saved_path = os.path.join(folder, "gemini_generated.png")
                    img.save(saved_path)
                    break
        if img:
            # Encode to base64 for inline display
            buffer = BytesIO()
            img.save(buffer, format="PNG")
            encoded = base64.b64encode(buffer.getvalue()).decode()
            return f"""
            ✅ <b>Gemini image generated successfully!</b><br>
            <img src='data:image/png;base64,{encoded}' width='400'><br>
            <small>📁 Saved at: {saved_path}</small>
            """
    except Exception as e:
        print(f"[Gemini Image Error] {e}")

    # --- Fallback to Hugging Face ---
    try:
        img = hf_client.text_to_image(prompt)
        saved_path = os.path.join(folder, "hf_generated.png")
        img.save(saved_path)

        # Encode image to base64
        buffer = BytesIO()
        img.save(buffer, format="PNG")
        encoded = base64.b64encode(buffer.getvalue()).decode()

        return f"""
        ✅ <b>Hugging Face image generated successfully!</b><br>
        <img src='data:image/png;base64,{encoded}' width='400'><br>
        <small>📁 Saved at: {saved_path}</small>
        """
    except Exception as e:
        return f"❌ Image generation failed: {e}"


# ============================ #
# Step 10: File & Image Processing Agents
# ============================ #
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from transformers import BlipProcessor, BlipForConditionalGeneration
import easyocr
import numpy as np

# ---------- Utility: Text Extraction ----------
def extract_text_from_file(uploaded_file):
    """Extracts text from TXT, PDF, or DOCX file."""
    name = uploaded_file.name.lower()
    if name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8", errors="ignore")
    elif name.endswith(".pdf"):
        reader = PdfReader(uploaded_file)
        return " ".join(page.extract_text() or "" for page in reader.pages)
    elif name.endswith(".docx"):
        doc = DocxDocument(uploaded_file)
        return " ".join(p.text for p in doc.paragraphs)
    return ""

# ---------- Utility: Text Summarization ----------
def summarize_text_with_llm(text):
    """Summarizes extracted text using Gemini."""
    if not text.strip():
        return "⚠️ No readable text found."
    prompt = f"Summarize this medical or research-related content in 5 concise bullet points:\n\n{text[:4000]}"
    try:
        result = llm.invoke(prompt)
        return getattr(result, "content", str(result))
    except Exception as e:
        return f"❌ Error summarizing text: {e}"

# ---------- Cached Model Loaders ----------
@st.cache_resource
def load_blip_model():
    """Load BLIP model for image captioning."""
    processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-large")
    model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-large")
    return processor, model

@st.cache_resource
def load_easyocr_reader():
    """Load EasyOCR model for text recognition."""
    return easyocr.Reader(["en"], gpu=False)

# ---------- Image Analyzer ----------
def describe_image_with_blip_and_ocr(image):
    """Uses BLIP for captioning and EasyOCR for text detection in an image."""
    try:
        processor, model = load_blip_model()
        inputs = processor(image, return_tensors="pt")
        out = model.generate(**inputs, max_new_tokens=80)
        caption = processor.decode(out[0], skip_special_tokens=True)

        reader = load_easyocr_reader()
        result = reader.readtext(np.array(image))
        extracted_text = " ".join([res[1] for res in result if res[1].strip()])

        response = f"🖼️ **BLIP Caption:** {caption}"
        if extracted_text:
            text_summary = summarize_text_with_llm(extracted_text)
            response += f"\n\n🧾 **Detected Text:** {extracted_text}\n\n🧠 **Text Summary:** {text_summary}"
        else:
            response += "\n\n🔍 No visible text detected in the image."

        return response
    except Exception as e:
        return f"❌ Error analyzing image: {e}"

# ---------------- Agent Icons ----------------
AGENT_ICONS = {
    "Drug Info Agent": "https://cdn-icons-png.flaticon.com/512/4320/4320365.png",
    "BMI Agent": "https://cdn-icons-png.flaticon.com/512/10476/10476452.png",
    "Diagnosis Agent": "https://cdn-icons-png.flaticon.com/512/5871/5871554.png",
    "Lifestyle Agent": "https://cdn-icons-png.flaticon.com/512/4310/4310163.png",
    "Research Agent": "https://cdn-icons-png.flaticon.com/512/3077/3077325.png",
    "Image Agent": "https://img.icons8.com/color/48/picture.png",
    "Image Caption & Text Agent": "https://img.icons8.com/color/48/image.png",
    "File Agent": "https://img.icons8.com/color/48/opened-folder.png",
    "Non-Medical": "https://img.icons8.com/color/48/error.png",
    "General": "https://cdn-icons-png.flaticon.com/512/14958/14958350.png"
}


# ============================ #
# Step 13: Streamlit Interface #
# ============================ #

# Preserve bold formatting in chat
def convert_markdown_to_html(text):
    """Preserve bold (<b>) and italic (<i>) formatting for black bold output."""
    text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*(.*?)\*", r"<i>\1</i>", text)
    return text


# ---------------- Page Config ----------------
st.set_page_config(page_title="🩺 MEDIBOT", page_icon="💊", layout="wide")

# ---------------- Sidebar ----------------
st.sidebar.markdown(
    """
    <style>
    /* ---------- Sidebar Quote Styling ---------- */
    .sidebar-quote {
        background: rgba(255, 255, 255, 0.08);
        border-radius: 12px;
        padding: 14px 16px;
        text-align: center;
        color: #e6f2ff;
        font-style: italic;
        font-size: 14px;
        line-height: 1.6;
        box-shadow: 0 0 15px rgba(0, 150, 255, 0.15);
        margin-top: 15px;
        transition: transform 0.4s ease, box-shadow 0.4s ease, background 0.4s ease;
        animation: heartbeat 3s ease-in-out infinite;
    }
    .sidebar-quote:hover {
        transform: translateY(-5px) scale(1.02);
        background: rgba(100, 180, 255, 0.15);
        box-shadow: 0 0 30px rgba(37, 117, 252, 0.4);
    }
    @keyframes heartbeat {
        0%, 100% { transform: scale(1); }
        50% { transform: scale(1.03); }
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Animation Image (Top)
st.sidebar.image("https://cdn-icons-png.flaticon.com/512/3558/3558977.png", width=120)

# Healing Quote (Now directly under animation)
st.sidebar.markdown(
    """
    <div class="sidebar-quote">
        💙 “Healing is not just in medicine — it’s in listening, empathy, and care.”  
        <br><span style="font-size:13px; color:#aee1ff;">— MEDIBOT, Guided by Compassion & Intelligence</span>
    </div>
    """,
    unsafe_allow_html=True
)

st.sidebar.markdown("---")

# ---------- FILE UPLOAD ----------
st.sidebar.markdown("<div class='upload-box'>📂 <b>Browse Files</b></div>", unsafe_allow_html=True)
uploaded_file = st.sidebar.file_uploader("Upload a document", type=["txt", "pdf", "docx"], label_visibility="collapsed")

if uploaded_file:
    progress = st.sidebar.progress(0)
    for pct in range(0, 101, 10):
        time.sleep(0.05)
        progress.progress(pct)
    st.sidebar.success(f"✅ {uploaded_file.name} uploaded successfully!")
    progress.empty()

    text = extract_text_from_file(uploaded_file)
    summary = summarize_text_with_llm(text)

    summary_html = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", summary)
    summary_html = re.sub(r"\*(.*?)\*", r"<i>\1</i>", summary_html)

    icon_url = AGENT_ICONS["File Agent"]
    st.session_state.messages.append({
        "role": "assistant",
        "content": f"<b>📄 Summary for file:</b> `{uploaded_file.name}`<br><br>{summary_html}",
        "icon": icon_url
    })

# ---------- IMAGE UPLOAD ----------
st.sidebar.markdown("<div class='upload-box'>🖼️ <b>Browse Images</b></div>", unsafe_allow_html=True)
uploaded_image = st.sidebar.file_uploader("Upload an image", type=["jpg", "jpeg", "png"], label_visibility="collapsed")

if uploaded_image:
    progress = st.sidebar.progress(0)
    for pct in range(0, 101, 10):
        time.sleep(0.05)
        progress.progress(pct)
    st.sidebar.success(f"✅ {uploaded_image.name} uploaded successfully!")
    progress.empty()

    image = Image.open(uploaded_image).convert("RGB")
    image_summary = describe_image_with_blip_and_ocr(image)

    image_summary_html = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", image_summary)
    image_summary_html = re.sub(r"\*(.*?)\*", r"<i>\1</i>", image_summary_html)

    icon_url = AGENT_ICONS["Image Caption & Text Agent"]
    st.session_state.messages.append({
        "role": "assistant",
        "content": f"<b>🖼️ Analysis for image:</b> `{uploaded_image.name}`<br><br>{image_summary_html}",
        "icon": icon_url
    })

st.sidebar.markdown("---")
if st.sidebar.button("🧹 Clear Chat"):
    st.session_state.clear()

# ---------------- Title ----------------
st.markdown("""

<style>
.title-container {
    text-align: center;
    margin-top: -15px;
    margin-bottom: 25px;
    animation: fadeIn 1.2s ease-in-out;
}
.main-title {
    font-size: 4.4rem;
    font-weight: 900;
    background: linear-gradient(90deg, #6a11cb, #2575fc);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    letter-spacing: 2px;
    text-shadow: 0px 0px 25px rgba(37,117,252,0.6);
    transition: transform 0.6s ease, text-shadow 0.6s ease;
    animation: glow 2.5s ease-in-out infinite alternate;
}
.main-title:hover {
    transform: translateY(-8px) scale(1.05);
    text-shadow: 0 0 40px rgba(106,17,203,0.8), 0 0 60px rgba(37,117,252,0.8);
}
.subtitle {
    font-size: 1rem;
    font-weight: 500;
    color: #555;
    margin-top: 6px;
    letter-spacing: 0.5px;
    transition: transform 0.4s ease, color 0.4s ease;
}
.subtitle:hover {
    transform: translateY(-4px);
    color: #2575fc;
}
@keyframes glow {
    from { text-shadow: 0 0 15px rgba(37,117,252,0.3); }
    to { text-shadow: 0 0 40px rgba(106,17,203,0.7); }
}
/* Sidebar Image Animation */
section[data-testid="stSidebar"] img {
    display: block;
    margin: 0 auto 10px auto;
    border-radius: 50%;
    box-shadow: 0 0 25px rgba(255,255,255,0.25);
    transition: all 0.4s ease-in-out;
    animation: fadeIn 1.2s ease-in-out;
}
section[data-testid="stSidebar"] img:hover {
    transform: scale(1.1);
    box-shadow: 0 0 40px rgba(100,150,255,0.8);
}
/* Upload Boxes */
.upload-box {
    background: linear-gradient(90deg, #6a11cb 0%, #2575fc 100%);
    color: white !important;
    text-align: center;
    padding: 10px;
    margin: 10px 0;
    border-radius: 12px;
    font-weight: bold;
    transition: transform 0.3s ease, box-shadow 0.3s ease;
    cursor: pointer;
    box-shadow: 0 3px 10px rgba(0, 0, 0, 0.2);
}
.upload-box:hover {
    transform: translateY(-4px);
    box-shadow: 0 5px 20px rgba(37, 117, 252, 0.5);
}
/* Clear Chat Button */
div[data-testid="stButton"] > button {
    background: linear-gradient(90deg, #6a11cb 0%, #2575fc 100%) !important;
    color: white !important;
    border: none !important;
    border-radius: 12px !important;
    font-weight: bold !important;
    transition: transform 0.3s ease, box-shadow 0.3s ease !important;
}
div[data-testid="stButton"] > button:hover {
    transform: translateY(-3px);
    box-shadow: 0 5px 20px rgba(37, 117, 252, 0.5) !important;
}
@keyframes fadeIn {
    from { opacity: 0; transform: scale(0.9); }
    to { opacity: 1; transform: scale(1); }
}
b { color: white; }
</style>
""", unsafe_allow_html=True)


# ---------------- Title ----------------
st.markdown("""
<div class="title-container">
    <h1 class="main-title">🩺 MEDIBOT</h1>
    <h2 class="subtitle">AI Powered Medical Assistant 💬</h2>
</div>

<style>

/* ---------- Centered Title Styling ---------- */
.title-container {
    text-align: center;
    margin-top: -15px;
    margin-bottom: 25px;
    animation: fadeIn 1.2s ease-in-out;
}

.main-title {
    font-size: 4.4rem;
    font-weight: 900;
    background: linear-gradient(90deg, #6a11cb, #2575fc);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    letter-spacing: 2px;
    text-shadow: 0px 0px 25px rgba(37,117,252,0.6);
    transition: transform 0.6s ease, text-shadow 0.6s ease;
    animation: glow 2.5s ease-in-out infinite alternate;
}

.main-title:hover {
    transform: translateY(-8px) scale(1.05);
    text-shadow: 0 0 40px rgba(106,17,203,0.8), 0 0 60px rgba(37,117,252,0.8);
}

.subtitle {
    font-size: 1rem;
    font-weight: 500;
    color: #555;
    margin-top: 6px;
    letter-spacing: 0.5px;
    transition: transform 0.4s ease, color 0.4s ease;
}

.subtitle:hover {
    transform: translateY(-4px);
    color: #2575fc;
}

/* ---------- Glowing Title Animation ---------- */
@keyframes glow {
    from {
        text-shadow: 0 0 15px rgba(37,117,252,0.3), 0 0 25px rgba(37,117,252,0.3);
    }
    to {
        text-shadow: 0 0 40px rgba(106,17,203,0.7), 0 0 60px rgba(37,117,252,0.7);
    }
}

/* ---------- SIDEBAR IMAGE ANIMATION EFFECT ---------- */
section[data-testid="stSidebar"] img {
    display: block;
    margin: 0 auto 25px auto;
    border-radius: 50%;
    box-shadow: 0 0 25px rgba(255,255,255,0.25);
    transition: all 0.4s ease-in-out;
    animation: fadeIn 1.2s ease-in-out;
}
section[data-testid="stSidebar"] img:hover {
    transform: scale(1.1);
    box-shadow: 0 0 40px rgba(100,150,255,0.8);
}

/* ---------- Floating Upload Boxes ---------- */
.upload-box {
    background: linear-gradient(90deg, #6a11cb 0%, #2575fc 100%);
    color: white !important;
    text-align: center;
    padding: 10px;
    margin: 10px 0;
    border-radius: 12px;
    font-weight: bold;
    transition: transform 0.3s ease, box-shadow 0.3s ease;
    cursor: pointer;
    box-shadow: 0 3px 10px rgba(0, 0, 0, 0.2);
}
.upload-box:hover {
    transform: translateY(-4px);
    box-shadow: 0 5px 20px rgba(37, 117, 252, 0.5);
}

/* ---------- Clear Chat Button ---------- */
div[data-testid="stButton"] > button {
    background: linear-gradient(90deg, #6a11cb 0%, #2575fc 100%) !important;
    color: white !important;
    border: none !important;
    border-radius: 12px !important;
    font-weight: bold !important;
    transition: transform 0.3s ease, box-shadow 0.3s ease !important;
}
div[data-testid="stButton"] > button:hover {
    transform: translateY(-3px);
    box-shadow: 0 5px 20px rgba(37, 117, 252, 0.5) !important;
}

/* ---------- Smooth fade-in animation ---------- */
@keyframes fadeIn {
    from { opacity: 0; transform: scale(0.9); }
    to { opacity: 1; transform: scale(1); }
}

/* ---------- Chat Bubbles ---------- */
.flex-container {
    display: flex;
    align-items: flex-end;
    margin: 10px 0;
    width: 100%;
}
.flex-start { justify-content: flex-start; }
.flex-end { justify-content: flex-end; }

.chat-bubble {
    max-width: 70%;
    padding: 12px 16px;
    border-radius: 18px;
    line-height: 1.4;
    word-wrap: break-word;
    font-size: 15px;
    color: white;
    transition: transform 0.3s ease, box-shadow 0.3s ease;
}

/* ---------- User Bubble ---------- */
.user-bubble {
    background: linear-gradient(90deg, #6a11cb 0%, #2575fc 100%);
    border-bottom-right-radius: 4px;
    box-shadow: 0 4px 12px rgba(106,17,203,0.3);
}
.user-bubble:hover {
    transform: translateY(-3px);
    box-shadow: 0 6px 18px rgba(37,117,252,0.6);
}

/* ---------- Assistant Bubble ---------- */
.assistant-bubble {
    background: linear-gradient(90deg, #2575fc 0%, #6a11cb 100%);
    border-bottom-left-radius: 4px;
    box-shadow: 0 4px 12px rgba(37,117,252,0.3);
}
.assistant-bubble:hover {
    transform: translateY(-3px);
    box-shadow: 0 6px 18px rgba(106,17,203,0.6);
}



/* ---------- User & Bot Icons ---------- */
.user-icon, .bot-icon {
    width: 40px;
    height: 40px;
    border-radius: 50%;
    background-size: cover;
    background-position: center;
}
.user-icon {
    background-image: url('https://img.icons8.com/color/48/user.png');
    margin-left: 8px;
}
.bot-icon {
    background-image: url('https://cdn-icons-png.flaticon.com/512/14958/14958350.png');
    margin-right: 8px;
}

b { color: white; }

</style>
""", unsafe_allow_html=True)

# ---------------- Session State ----------------
if "messages" not in st.session_state:
    st.session_state.messages = []


    # ✅ Sync LangChain memory after image analysis
    memory.chat_memory.messages = []
    for msg in st.session_state.messages:
        if msg["role"] == "user":
            memory.chat_memory.add_user_message(msg["content"])
        elif msg["role"] == "assistant":
            memory.chat_memory.add_ai_message(msg["content"])
            
            

# ---------------- Chat Input ----------------
user_input = st.chat_input("Type your medical question...")

# ---------------- Routing & Logic ----------------
if user_input:
    # Save user message (persistent)
    st.session_state.messages.append({"role": "user", "content": user_input})
    q_lower = user_input.lower().strip()
    agent_name, response = "Unknown", ""

    # ---------------- Exit Handling ----------------
    if q_lower in ["quit", "exit", "bye", "goodbye", "stop", "end"]:
        farewell_msg = (
            "🩺 **Thank you!** Take care of your health and well-being. "
            "If you have more medical questions later, feel free to return anytime! 🙂"
        )
        st.session_state.messages.append({
            "role": "assistant",
            "content": farewell_msg,
            "icon": AGENT_ICONS.get("General")
        })

        memory.chat_memory.clear()
        st.session_state.clear()
        st.rerun()

    # ---------------- Agent Routing ----------------
    elif any(w in q_lower for w in ["research", "study", "journal", "paper", "article"]):
        agent_name, response = "Research Agent", research_agent(user_input, llm=llm)

    elif is_symptom_query(user_input):
        agent_name, response = "Diagnosis Agent", diagnosis_agent(user_input)

    elif is_drug_query(user_input):
        agent_name, response = "Drug Info Agent", get_drug_info(is_drug_query(user_input))

    elif is_bmi_query(user_input):
        agent_name, response = "BMI Agent", bmi_agent(user_input)

    elif any(w in q_lower for w in ["diet", "exercise", "lifestyle", "fitness", "nutrition", "wellness", "exercises"]):
        agent_name, response = "Lifestyle Agent", lifestyle_agent(user_input, llm=llm)

    elif any(w in q_lower for w in ["sketch", "diagram", "draw", "image", "illustration", "visualize"]):
        agent_name, response = "Image Agent", image_agent(user_input)

    elif not is_medical_question(user_input):
        agent_name, response = "Non-Medical", "⚠️ This question is outside the medical domain."

    # ---------------- Follow-Up Detection + RAG ----------------
    elif detect_followup(user_input):
        # ✅ Follow-up detected: reuse previous user question + assistant answer
        last_user_query = ""
        last_ai_answer = ""
        for msg in reversed(st.session_state.messages):
            if not last_user_query and msg["role"] == "user":
                last_user_query = msg["content"]
            elif not last_ai_answer and msg["role"] == "assistant":
                last_ai_answer = msg["content"]
            if last_user_query and last_ai_answer:
                break

        followup_prompt = (
            f"The user said '{user_input}'. "
            f"This is a follow-up referring to the previous topic. "
            f"The previous user question was: '{last_user_query}'. "
            f"The assistant previously explained: '{last_ai_answer}'. "
            f"Continue the same topic, providing further details, applications, or related insights."
        )

        try:
            result = medical_chain({
                "question": followup_prompt,
                "chat_history": get_combined_history(
                    memory.load_memory_variables({}).get("chat_history", [])
                )
            })
            response = result["answer"].strip()
            agent_name = "RAG Chain (Follow-Up)"
        except Exception as e:
            response = f"❌ Error during follow-up handling: {e}"

    else:
        # ---------------- Normal RAG Retrieval ---------------- #
        docs = retriever.get_relevant_documents(user_input)
        context_text = " ".join([doc.page_content for doc in docs]) if docs else ""
        full_history = memory.load_memory_variables({}).get("chat_history", [])
        update_summary_memory(full_history)
        chat_history_str = get_combined_history(full_history)

        try:
            result = medical_chain({
                "question": user_input,
                "chat_history": chat_history_str
            })
            response = result["answer"].strip()
            agent_name = "RAG Chain"
        except Exception as e:
            response = f"❌ Error generating answer: {e}"

    # ---------------- Sync LangChain memory with Streamlit session ----------------
    memory.chat_memory.messages = []
    for msg in st.session_state.messages:
        if msg["role"] == "user":
            memory.chat_memory.add_user_message(msg["content"])
        elif msg["role"] == "assistant":
            memory.chat_memory.add_ai_message(msg["content"])

    # ---------------- Format & Append ----------------
    formatted_response = convert_markdown_to_html(str(response))
    icon_url = AGENT_ICONS.get(agent_name, AGENT_ICONS["General"])

    if agent_name in AGENT_ICONS and agent_name != "General":
        st.session_state.messages.append({
            "role": "assistant",
            "content": f"**{agent_name} Activated**",
            "icon": icon_url
        })

    st.session_state.messages.append({
        "role": "assistant",
        "content": formatted_response,
        "icon": icon_url
    })

    # Re-sync to store the latest reply in memory
    memory.chat_memory.messages = []
    for msg in st.session_state.messages:
        if msg["role"] == "user":
            memory.chat_memory.add_user_message(msg["content"])
        elif msg["role"] == "assistant":
            memory.chat_memory.add_ai_message(msg["content"])

    st.rerun()

# ---------------- Display Function ----------------
def display_message(msg):
    clean_content = msg["content"]
    if msg["role"] == "user":
        st.markdown(f"""
        <div class="flex-container flex-end">
            <div class='chat-bubble assistant-bubble'>{clean_content}</div>
            <div class='user-icon'></div>
        </div>
        """, unsafe_allow_html=True)
    else:
        icon_url = msg.get("icon", AGENT_ICONS["General"])
        st.markdown(f"""
        <div class="flex-container flex-start">
            <div class='bot-icon' style="background-image:url('{icon_url}')"></div>
            <div class='chat-bubble assistant-bubble'>{clean_content.replace(chr(10), '<br>')}</div>
        </div>
        """, unsafe_allow_html=True)

# ---------------- Display All Messages ----------------
for i, msg in enumerate(st.session_state.messages):
    if i == len(st.session_state.messages) - 1 and msg["role"] == "assistant":
        continue
    display_message(msg)

# ---------------- Typing Animation ----------------
if st.session_state.messages and st.session_state.messages[-1]["role"] == "assistant":
    final_msg = st.session_state.messages[-1]["content"]
    icon_url = st.session_state.messages[-1].get("icon", AGENT_ICONS["General"])
    placeholder = st.empty()
    lines = final_msg.split("\n")
    full_text = ""
    for line in lines:
        full_text += line + "\n"
        placeholder.markdown(f"""
        <div class="flex-container flex-start">
            <div class='bot-icon' style="background-image:url('{icon_url}')"></div>
            <div class='chat-bubble assistant-bubble'>{full_text.replace(chr(10), '<br>')}</div>
        </div>
        """, unsafe_allow_html=True)
        time.sleep(0.3)
